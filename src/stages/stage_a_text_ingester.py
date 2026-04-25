"""
DIAS Stage A - Text Ingester

Advanced text extraction and intelligent chunking system for PDF, EPUB, and DOCX files.
Uses dual parsing strategy for maximum compatibility and intelligent text chunking
that preserves narrative coherence.
"""

import os
import re
import json
import logging
from pathlib import Path
from typing import List, Dict, Optional, Union, Tuple, Any
from dataclasses import dataclass
from enum import Enum
from uuid import uuid4

import fitz  # PyMuPDF
import pdfplumber
from bs4 import BeautifulSoup
import ebooklib
from ebooklib import epub
from docx import Document
import nltk
import spacy
from tqdm import tqdm

import sys
from pathlib import Path

# Aggiungi il path root al Python path per trovare il modulo 'src'
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from src.common.persistence import DiasPersistence
from src.common.base_stage import BaseStage
from src.common.redis_factory import get_redis_client
from src.common.config import get_config
from src.common.models import IngestionBlock, BookMetadata
from src.common.logging_setup import setup_logging

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

# Load spaCy model for intelligent chunking
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    logging.warning("spaCy English model not found. Install with: python -m spacy download en_core_web_sm")
    nlp = None


class FileFormat(Enum):
    """Supported file formats for text extraction."""
    PDF = "pdf"
    EPUB = "epub" 
    DOCX = "docx"
    TXT = "txt"
    UNKNOWN = "unknown"


@dataclass
class TextChunk:
    """Represents a chunk of text with metadata."""
    text: str
    chunk_id: int
    word_count: int
    start_index: int
    end_index: int
    chapter_hint: Optional[str] = None
    paragraph_boundaries: List[int] = None
    
    def __post_init__(self):
        if self.paragraph_boundaries is None:
            self.paragraph_boundaries = []


class TextIngester(BaseStage):
    """
    Stage A: Text Ingester
    
    Extracts text from PDF, EPUB, and DOCX files and intelligently chunks them
    into 2500-word segments with 500-word overlap while preserving narrative
    coherence by respecting paragraph and chapter boundaries.
    
    Attributes:
        supported_formats: List of supported file formats
        chunk_size: Target size for text chunks (default: 2500 words)
        overlap_size: Overlap size between chunks (default: 500 words)
        min_chunk_size: Minimum chunk size (default: 1500 words)
        max_chunk_size: Maximum chunk size (default: 3500 words)
    """
    
    def __init__(self, redis_client=None, config=None):
        """
        Initialize the Text Ingester.
        
        Args:
            redis_client: Redis client for queue operations (optional, uses factory if None)
            config: Configuration object (optional, uses get_config() if None)
        """
        # Get config first (needed for logging)
        self.config = config or get_config()
        
        # Setup logger (required before BaseStage init)
        self.logger = setup_logging(
            "stage_a_text_ingester",
            level=self.config.logging.level,
            log_file=self.config.logging.file,
        )
        
        # Get Redis client from factory if not provided
        if redis_client is None:
            redis_client = get_redis_client(logger=self.logger)
            self.logger.info("Using config-driven Redis client")
        
        super().__init__(
            stage_name="text_ingester",
            stage_number=1,
            input_queue="dias:q:0:upload",  # Keep upload as is for now or use config if exists
            output_queue=None, # Set below
            config=config,
            redis_client=redis_client
        )
        self.output_queue = self.config.queues.ingestion
        
        # Initialize persistence manager
        # self.persistence = DiasPersistence() # Decommissioned global persistence
        self.persistence = DiasPersistence() # Kept for utility (normalize_id) but will be overridden in process
        
        # Temporary storage for testing
        self._temp_blocks = []
        
        # Configuration
        self.supported_formats = [FileFormat.PDF, FileFormat.EPUB, FileFormat.DOCX]
        self.chunk_size = 2500  # Target words per chunk
        self.overlap_size = 500  # Words overlap between chunks
        self.min_chunk_size = 1500  # Minimum words per chunk
        self.max_chunk_size = 3500  # Maximum words per chunk
        
        # Logger
        self.logger = setup_logging(f"{self.__class__.__name__}")
        
        self.logger.info("TextIngester initialized with intelligent chunking")
    
    def detect_file_format(self, file_path: Union[str, Path]) -> FileFormat:
        """
        Detect file format based on extension and content inspection.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Detected file format
        """
        file_path = Path(file_path)
        
        # Check extension first
        ext = file_path.suffix.lower()
        format_map = {
            '.pdf': FileFormat.PDF,
            '.epub': FileFormat.EPUB,
            '.docx': FileFormat.DOCX,
            '.txt': FileFormat.TXT
        }
        
        detected_format = format_map.get(ext, FileFormat.UNKNOWN)
        
        # Content inspection for additional validation
        if detected_format != FileFormat.UNKNOWN:
            try:
                with open(file_path, 'rb') as f:
                    header = f.read(100)
                    
                if detected_format == FileFormat.PDF and not header.startswith(b'%PDF'):
                    self.logger.warning(f"File {file_path} has PDF extension but invalid header")
                    detected_format = FileFormat.UNKNOWN
                    
                elif detected_format == FileFormat.EPUB:
                    # EPUB is essentially a ZIP file
                    if not header.startswith(b'PK'):
                        self.logger.warning(f"File {file_path} has EPUB extension but invalid ZIP header")
                        detected_format = FileFormat.UNKNOWN
                        
            except Exception as e:
                self.logger.error(f"Error inspecting file {file_path}: {e}")
                detected_format = FileFormat.UNKNOWN
        
        return detected_format
    
    def extract_text_from_pdf(self, file_path: Union[str, Path]) -> str:
        """
        Extract text from PDF using dual parsing strategy.
        
        Primary: PyMuPDF for performance and complex PDF support
        Fallback: pdfplumber for open-source compatibility
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
            
        Raises:
            ValueError: If PDF parsing fails
        """
        file_path = Path(file_path)
        text_parts = []
        
        self.logger.info(f"Extracting text from PDF: {file_path}")
        
        try:
            # Primary strategy: PyMuPDF
            with fitz.open(file_path) as doc:
                total_pages = len(doc)
                self.logger.info(f"PDF has {total_pages} pages")
                
                for page_num in tqdm(range(total_pages), desc="Processing PDF pages"):
                    page = doc.load_page(page_num)
                    page_text = page.get_text()
                    
                    if page_text.strip():
                        text_parts.append(page_text)
                    else:
                        self.logger.warning(f"Page {page_num + 1} appears to be empty or image-only")
                        
                        # Try OCR for image-only pages (placeholder for future enhancement)
                        # This would require additional dependencies like pytesseract
                        
            extracted_text = '\n\n'.join(text_parts)
            
            if not extracted_text.strip():
                raise ValueError("PyMuPDF extracted no text, trying fallback")
                
            self.logger.info(f"Successfully extracted {len(extracted_text)} characters using PyMuPDF")
            return extracted_text
            
        except Exception as primary_error:
            self.logger.warning(f"PyMuPDF failed: {primary_error}. Trying pdfplumber fallback.")
            
            try:
                # Fallback strategy: pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    total_pages = len(pdf.pages)
                    self.logger.info(f"PDF has {total_pages} pages (pdfplumber)")
                    
                    for page_num in tqdm(range(total_pages), desc="Processing PDF pages (fallback)"):
                        page = pdf.pages[page_num]
                        page_text = page.extract_text()
                        
                        if page_text and page_text.strip():
                            text_parts.append(page_text)
                        else:
                            self.logger.warning(f"Page {page_num + 1} appears to be empty")
                    
                    extracted_text = '\n\n'.join(text_parts)
                    
                    if not extracted_text.strip():
                        raise ValueError("pdfplumber also extracted no text")
                        
                    self.logger.info(f"Successfully extracted {len(extracted_text)} characters using pdfplumber")
                    return extracted_text
                    
            except Exception as fallback_error:
                error_msg = f"Both PDF parsers failed. PyMuPDF: {primary_error}, pdfplumber: {fallback_error}"
                self.logger.error(error_msg)
                raise ValueError(error_msg)
    
    def extract_text_from_epub(self, file_path: Union[str, Path]) -> str:
        """
        Extract text from EPUB file.
        
        Args:
            file_path: Path to EPUB file
            
        Returns:
            Extracted text
            
        Raises:
            ValueError: If EPUB parsing fails
        """
        file_path = Path(file_path)
        text_parts = []
        
        self.logger.info(f"Extracting text from EPUB: {file_path}")
        
        try:
            book = epub.read_epub(file_path)
            
            # Get book metadata
            title = book.get_metadata('DC', 'title')
            author = book.get_metadata('DC', 'creator')
            
            if title:
                self.logger.info(f"EPUB Title: {title[0][0] if title else 'Unknown'}")
            if author:
                self.logger.info(f"EPUB Author: {author[0][0] if author else 'Unknown'}")
            
            # Extract text from all documents
            items = list(book.get_items_of_type(ebooklib.ITEM_DOCUMENT))
            self.logger.info(f"EPUB has {len(items)} content documents")
            
            for item in tqdm(items, desc="Processing EPUB documents"):
                try:
                    # Parse HTML content
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    
                    # Remove script and style elements
                    for script in soup(["script", "style"]):
                        script.decompose()
                    
                    # Extract text
                    text = soup.get_text()
                    
                    # Clean up whitespace
                    lines = (line.strip() for line in text.splitlines())
                    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                    text = ' '.join(chunk for chunk in chunks if chunk)
                    
                    if text.strip():
                        text_parts.append(text)
                        
                except Exception as e:
                    self.logger.warning(f"Error processing EPUB item {item.get_id()}: {e}")
                    continue
            
            extracted_text = '\n\n'.join(text_parts)
            
            if not extracted_text.strip():
                raise ValueError("No text extracted from EPUB")
                
            self.logger.info(f"Successfully extracted {len(extracted_text)} characters from EPUB")
            return extracted_text
            
        except Exception as e:
            error_msg = f"EPUB extraction failed: {e}"
            self.logger.error(error_msg)
            raise ValueError(error_msg)
    
    def extract_text_from_docx(self, file_path: Union[str, Path]) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
            
        Raises:
            ValueError: If DOCX parsing fails
        """
        file_path = Path(file_path)
        text_parts = []
        
        self.logger.info(f"Extracting text from DOCX: {file_path}")
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract text from tables (if any)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(' | '.join(row_text))
            
            extracted_text = '\n\n'.join(paragraphs)
            
            if not extracted_text.strip():
                raise ValueError("No text extracted from DOCX")
                
            self.logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX")
            return extracted_text
            
        except Exception as e:
            error_msg = f"DOCX extraction failed: {e}"
            self.logger.error(error_msg)
            raise ValueError(error_msg)
    
    def extract_text(self, file_path: Union[str, Path]) -> str:
        """
        Extract text from file based on detected format.
        
        Args:
            file_path: Path to input file
            
        Returns:
            Extracted text
            
        Raises:
            ValueError: If file format is unsupported or extraction fails
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise ValueError(f"File not found: {file_path}")
        
        format_type = self.detect_file_format(file_path)
        
        if format_type == FileFormat.PDF:
            return self.extract_text_from_pdf(file_path)
        elif format_type == FileFormat.EPUB:
            return self.extract_text_from_epub(file_path)
        elif format_type == FileFormat.DOCX:
            return self.extract_text_from_docx(file_path)
        elif format_type == FileFormat.TXT:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
    
    def find_sentence_boundaries(self, text: str) -> List[int]:
        """
        Find sentence boundaries using NLTK and spaCy for intelligent chunking.
        
        Args:
            text: Input text
            
        Returns:
            List of sentence boundary indices
        """
        boundaries = []
        
        try:
            # Use spaCy if available for better accuracy
            if nlp:
                doc = nlp(text)
                for sent in doc.sents:
                    boundaries.append(sent.end_char)
            else:
                # Fallback to NLTK
                sentences = nltk.sent_tokenize(text)
                current_pos = 0
                for sentence in sentences:
                    current_pos = text.find(sentence, current_pos) + len(sentence)
                    boundaries.append(current_pos)
                    
        except Exception as e:
            self.logger.warning(f"Sentence boundary detection failed: {e}")
            # Fallback to simple punctuation-based detection
            import re
            boundaries = [m.end() for m in re.finditer(r'[.!?]+\s+', text)]
        
        return boundaries
    
    def find_paragraph_boundaries(self, text: str) -> List[int]:
        """
        Find paragraph boundaries for intelligent chunking.
        
        Args:
            text: Input text
            
        Returns:
            List of paragraph boundary indices
        """
        # Look for double newlines as paragraph separators
        import re
        boundaries = []
        
        # Find double newlines (paragraph separators)
        for match in re.finditer(r'\n\s*\n', text):
            boundaries.append(match.end())
        
        # Also consider single newlines if they're followed by indentation
        for match in re.finditer(r'\n(?=\s)', text):
            boundaries.append(match.end())
        
        return sorted(set(boundaries))
    
    def intelligent_chunk_text(self, text: str, book_id: str) -> List[TextChunk]:
        """
        Intelligently chunk text while preserving narrative coherence.
        
        Uses multiple strategies:
        1. Respect paragraph boundaries
        2. Respect sentence boundaries  
        3. Look for chapter/section markers
        4. Maintain target chunk size with overlap
        
        Args:
            text: Input text to chunk
            book_id: Book identifier for metadata
            
        Returns:
            List of TextChunk objects
        """
        self.logger.info(f"Intelligently chunking text for book {book_id}")
        
        if not text.strip():
            self.logger.warning("Empty text provided for chunking")
            return []
        
        # Find natural boundaries
        paragraph_boundaries = self.find_paragraph_boundaries(text)
        sentence_boundaries = self.find_sentence_boundaries(text)
        
        # Look for chapter markers
        chapter_pattern = re.compile(r'\n\s*(?:chapter|capitolo|section|parte)\s+\d+', re.IGNORECASE)
        chapter_matches = list(chapter_pattern.finditer(text))
        
        chunks = []
        current_pos = 0
        chunk_id = 0
        
        while current_pos < len(text):
            # Find optimal chunk end
            chunk_end = self.find_optimal_chunk_end(
                text, current_pos, paragraph_boundaries, 
                sentence_boundaries, chapter_matches
            )
            
            if chunk_end <= current_pos:
                # Fallback: force chunk end
                chunk_end = min(current_pos + self.max_chunk_size * 5, len(text))  # Approx 5 chars per word
            
            chunk_text = text[current_pos:chunk_end].strip()
            
            if not chunk_text:
                break
            
            # Calculate word count (approximate)
            word_count = len(chunk_text.split())
            
            # Create chunk metadata
            chapter_hint = None
            for chapter_match in chapter_matches:
                if chapter_match.start() <= current_pos < chapter_match.end():
                    chapter_hint = chapter_match.group().strip()
                    break
            
            # Find paragraph boundaries within this chunk
            chunk_paragraphs = [p - current_pos for p in paragraph_boundaries 
                              if current_pos <= p < chunk_end]
            
            chunk = TextChunk(
                text=chunk_text,
                chunk_id=chunk_id,
                word_count=word_count,
                start_index=current_pos,
                end_index=chunk_end,
                chapter_hint=chapter_hint,
                paragraph_boundaries=chunk_paragraphs
            )
            
            chunks.append(chunk)
            self.logger.debug(f"Created chunk {chunk_id}: {word_count} words")
            
            # Move to next chunk with overlap
            overlap_start = max(0, chunk_end - self.overlap_size * 5)  # Approximate char count
            current_pos = overlap_start if overlap_start > current_pos else chunk_end
            chunk_id += 1
            
            # Safety check to prevent infinite loops
            if current_pos >= len(text) or chunk_id > 10000:
                break
        
        self.logger.info(f"Created {len(chunks)} chunks from text")
        return chunks
    
    def intelligent_micro_chunk(self, text: str, macro_index: int) -> List[Dict[str, Any]]:
        """
        Split a macro-chunk into micro-chunks (~300 words) for Stage C.
        Ensures sentence boundaries are respected.
        """
        target_micro_words = 300
        words = text.split()
        total_words = len(words)
        
        micro_chunks = []
        start_idx = 0
        micro_idx = 0
        
        while start_idx < total_words:
            # Target end
            end_idx = min(start_idx + target_micro_words, total_words)
            
            # Find a near period to avoid cutting sentences
            if end_idx < total_words:
                temp_text = " ".join(words[start_idx:end_idx+20]) # look ahead a bit
                last_period = temp_text.rfind('.', 0, len(" ".join(words[start_idx:end_idx])) + 50)
                if last_period != -1:
                    # Sync end_idx with character position of period
                    # Simple approximation: find how many words before that period
                    micro_text_raw = temp_text[:last_period+1]
                    end_idx = start_idx + len(micro_text_raw.split())
            
            micro_text = " ".join(words[start_idx:end_idx]).strip()
            if micro_text:
                micro_chunks.append({
                    "micro_index": micro_idx,
                    "text": micro_text,
                    "word_count": len(micro_text.split()),
                    "label": f"chunk-{macro_index:03d}-micro-{micro_idx:03d}"
                })
            
            start_idx = end_idx
            micro_idx += 1
            
        return micro_chunks

    def find_optimal_chunk_end(self, text: str, start_pos: int, 
                             paragraph_boundaries: List[int], 
                             sentence_boundaries: List[int],
                             chapter_matches: List) -> int:
        """
        Find optimal position to end a chunk based on multiple criteria.
        
        Args:
            text: Full text
            start_pos: Starting position for this chunk
            paragraph_boundaries: List of paragraph boundary positions
            sentence_boundaries: List of sentence boundary positions  
            chapter_matches: List of chapter match objects
            
        Returns:
            Optimal end position for the chunk
        """
        # Calculate target range
        min_end = start_pos + self.min_chunk_size * 5  # Approximate
        target_end = start_pos + self.chunk_size * 5  # Approximate
        max_end = start_pos + self.max_chunk_size * 5  # Approximate
        
        # Find best boundary within range
        best_end = max_end  # Default fallback
        
        # Look for paragraph boundaries in target range
        for boundary in paragraph_boundaries:
            if min_end <= boundary <= target_end:
                best_end = boundary
                break
            elif target_end < boundary <= max_end:
                best_end = boundary
                break
        
        # If no good paragraph boundary, look for sentence boundaries
        if best_end == max_end:
            for boundary in sentence_boundaries:
                if min_end <= boundary <= target_end:
                    best_end = boundary
                    break
                elif target_end < boundary <= max_end:
                    best_end = boundary
                    break
        
        # Check for chapter boundaries (never split across chapters)
        for chapter_match in chapter_matches:
            chapter_start = chapter_match.start()
            chapter_end = chapter_match.end()
            
            # If chunk would cross chapter boundary, end at chapter start
            if start_pos < chapter_start < best_end:
                best_end = chapter_start
                break
        
        # Ensure we don't go beyond text length
        best_end = min(best_end, len(text))

        return best_end

    def _load_or_build_chapter_boundaries(self, project_id: str, full_text: str) -> List[Dict]:
        """
        Return chapter boundary list from chapter_boundaries.json (cached) or build it
        from fingerprint.json using chapter_detector.

        Returns [] if no fingerprint is available or detection fails, which causes
        the caller to fall back to single-chapter (chapter_001) mode.
        """
        try:
            from src.common.chapter_detector import load_or_build_boundaries
        except ImportError:
            self.logger.warning("chapter_detector not available; falling back to single-chapter mode")
            return []

        project_root = self.persistence.project_root
        return load_or_build_boundaries(project_root, full_text)

    def _chunk_by_chapters(
        self,
        full_text: str,
        boundaries: List[Dict],
        project_id: str,
    ) -> List[tuple]:
        """
        Split full_text into sections defined by boundaries, then chunk each section.

        Returns a list of (chapter_id: str, chapter_number: int, TextChunk).
        Chunks are ordered: all chunks of chapter 1, then chapter 2, etc.
        Chunk IDs are sequential across the whole book (not per-chapter).
        """
        results = []
        global_chunk_idx = 0

        for i, boundary in enumerate(boundaries):
            start = boundary["start_char"]
            end = (
                boundaries[i + 1]["start_char"]
                if i + 1 < len(boundaries)
                else len(full_text)
            )
            chapter_text = full_text[start:end]
            if not chapter_text.strip():
                continue

            chapter_id = boundary["chapter_id"]
            chapter_number = boundary["chapter_number"]

            chapter_chunks = self.intelligent_chunk_text(chapter_text, project_id)
            for chunk in chapter_chunks:
                # Re-index chunk_id globally
                chunk.chunk_id = global_chunk_idx
                results.append((chapter_id, chapter_number, chunk))
                global_chunk_idx += 1

            self.logger.info(
                f"Chapter {chapter_id} ({boundary['name'][:40]}): "
                f"{len(chapter_chunks)} macro-chunks"
            )

        return results

    def process(self, message: dict) -> Optional[dict]:
        """
        Process a book file by extracting and chunking text.
        
        Args:
            message: Message containing book file information
            
        Returns:
            Processing result or None if failed
        """
        try:
            # [SANITY CHECK] Force the official sanitized ID as the only truth
            raw_id = message.get("project_id") or message.get("book_id") or "unknown"
            project_id = DiasPersistence.normalize_id(raw_id)
            
            # [BACKWARD COMPATIBILITY] Ensure legacy names are available for internal logic
            book_id = project_id
            clean_title = project_id
            
            title = message.get("title", project_id)
            file_path = message.get("file_path")

            if not project_id or project_id == "unknown":
                self.logger.error("No valid project_id provided in message")
                return None

            self.persistence = DiasPersistence(project_id=project_id)

            self.logger.info(f"=== Stage A Processing for project: {project_id} ===")
            
            saved_file_paths = {}
            
            # --- DETERMINISTIC SOURCE DETECTION ---
            # Use the provided file_path (which is the source or normalized text pointer)
            if not file_path:
                self.logger.error("No file_path provided in message")
                return None
                
            if os.path.isabs(file_path):
                source_path = Path(file_path)
            else:
                # Resolve relative to project root
                source_path = Path(self.persistence.project_root) / file_path

            if not source_path.exists():
                self.logger.error(f"❌ Source file not found: {source_path}")
                return None
            
            self.logger.info(f"📖 Stage A loading text from: {source_path}")

            # Extract text based on file format
            try:
                full_text = self.extract_text(source_path)
            except Exception as e:
                self.logger.error(f"Text extraction failed for {book_id}: {e}")
                return None
            
            if not full_text.strip():
                self.logger.error(f"No text extracted from {book_id}")
                return None
            
            # Build chapter boundaries (mathematical, no LLM calls)
            chapter_boundaries = self._load_or_build_chapter_boundaries(project_id, full_text)
            chapter_count = len(chapter_boundaries) if chapter_boundaries else 1

            # Create book metadata
            book_metadata = BookMetadata(
                book_id=book_id,
                title=message.get("title", "Unknown"),
                author=message.get("author", ""),
                word_count=len(full_text.split()),
                chapter_count=chapter_count,
                file_path=str(file_path),
                file_format=self.detect_file_format(file_path).value
            )

            # Chunk text — chapter-aware if boundaries available, else full-text fallback
            if chapter_boundaries:
                self.logger.info(f"Chapter-aware chunking: {chapter_count} chapters detected")
                chunks_with_chapters = self._chunk_by_chapters(full_text, chapter_boundaries, project_id)
            else:
                self.logger.info("No chapter boundaries found — chunking as single chapter (chapter_001)")
                raw_chunks = self.intelligent_chunk_text(full_text, project_id)
                chunks_with_chapters = [("chapter_001", 1, c) for c in raw_chunks]

            if not chunks_with_chapters:
                self.logger.error(f"Text chunking failed for {project_id}")
                return None

            total_chunks = len(chunks_with_chapters)
            self.logger.info(f"Created {total_chunks} total chunks for book {project_id}")

            # Convert to ingestion blocks — one per (chapter_id, chunk)
            ingestion_blocks = []
            for i, (ch_id, ch_num, chunk) in enumerate(chunks_with_chapters):
                block = IngestionBlock(
                    book_id=project_id,
                    chapter_id=ch_id,
                    chapter_number=ch_num,
                    block_id=str(uuid4()),
                    block_text=chunk.text,
                    word_count=chunk.word_count,
                    block_index=i,
                    total_blocks_in_chapter=total_chunks
                )
                ingestion_blocks.append(block)
            
            # Push chunks to output queue
            success_count = 0
            self._temp_blocks = list(ingestion_blocks)  # Save for testing
            # [FIX] Remosso clean_title manuale incoerente. Usiamo solo project_id.
            
            for block in ingestion_blocks:
                try:
                    # Prepara dati per salvataggio
                    block_data = {
                        "book_id": block.book_id,
                        "chapter_id": block.chapter_id,
                        "chapter_number": block.chapter_number,
                        "block_id": block.block_id,
                        "block_text": block.block_text,
                        "word_count": block.word_count,
                        "block_index": block.block_index,
                        "total_blocks_in_chapter": block.total_blocks_in_chapter,
                        "timestamp": block.timestamp.isoformat()
                    }
                    
                    # Prepara nome file leggibile: Titolo-chunk-000
                    chunk_label = f"chunk-{block.block_index:03d}"
                    
                    # Salva su disco e traccia percorso
                    filepath = self.persistence.save_stage_output(
                        stage="a",
                        data=block_data,
                        book_id=project_id,
                        block_id=chunk_label,
                        include_timestamp=False
                    )
                    saved_file_paths[block.block_id] = filepath
                    
                    self.logger.info(f"💾 Blocco salvato: {filepath}")
                    
                    # --- Scomposizione in MICRO-CHUNK per Stage C ---
                    micro_chunks = self.intelligent_micro_chunk(block.block_text, block.block_index)
                    for micro in micro_chunks:
                        micro_data = {
                            "book_id": block.book_id,
                            "chapter_id": block.chapter_id,      # propagate chapter
                            "chapter_number": block.chapter_number,
                            "macro_index": block.block_index,
                            "micro_index": micro["micro_index"],
                            "block_id": f"{block.book_id}-{micro['label']}",
                            "block_text": micro["text"],
                            "word_count": micro["word_count"],
                            "timestamp": block.timestamp.isoformat()
                        }
                        self.persistence.save_stage_output(
                            stage="a",
                            data=micro_data,
                            book_id=block.book_id,
                            block_id=micro["label"],
                            include_timestamp=False
                        )
                    self.logger.info(f"📦 Generati {len(micro_chunks)} micro-chunk per {chunk_label}")
                    # ----------------------------------------------------
                    
                except Exception as e:
                    self.logger.error(f"❌ Errore salvataggio blocco {block.block_id}: {e}")
                    continue
            
            # Ora processa i blocchi per Redis
            for block in ingestion_blocks:
                try:
                    # Convert metadata to serializable format
                    book_metadata_dict = book_metadata.model_dump()
                    book_metadata_dict['processing_timestamp'] = book_metadata_dict['processing_timestamp'].isoformat()
                    
                    message_data = {
                        "book_id": book_id,
                        "block_id": block.block_id,
                        "text": block.block_text,
                        "word_count": block.word_count,
                        "chapter_id": block.chapter_id,
                        "chapter_number": block.chapter_number,
                        "block_index": block.block_index,
                        "total_blocks_in_chapter": block.total_blocks_in_chapter,
                        "book_metadata": book_metadata_dict,
                        "clean_title": clean_title,
                        "chunk_label": f"chunk-{block.block_index:03d}",
                        "stage": "ingestion",
                        "timestamp": block.timestamp.isoformat(),
                        "file_path": saved_file_paths.get(block.block_id)  # Percorso file salvato
                    }
                    
                    self.redis.push_to_queue(self.output_queue, message_data)
                    success_count += 1
                    
                except Exception as e:
                    self.logger.error(f"Failed to push chunk {block.block_id}: {e}")
                    continue
            
            self.logger.info(f"Successfully pushed {success_count}/{len(ingestion_blocks)} chunks for book {book_id}")
            
            # Return processing summary
            # Salva checkpoint manualmente prima di uscire (BaseStage lo farebbe solo se restituissimo result)
            self.redis.set_checkpoint(book_id, self.stage_number)
            
            # Nuova logica Sprint 4: Aggiorna lo stato del progetto nel config.json
            try:
                self.persistence.update_project_config({"status": "ingested"})
                self.logger.info(f"📈 Project status updated to 'ingested' for {book_id}")
            except Exception as e:
                self.logger.warning(f"Failed to update project config for {book_id}: {e}")
            
            # Restituiamo None per evitare che BaseStage invii il messaggio di riepilogo alla coda dello Stage B
            return None
            
        except Exception as e:
            self.logger.error(f"Unexpected error processing message: {e}")
            return None
    
    def process_book_file(self, file_path: Union[str, Path], book_id: str, 
                         metadata: Optional[Dict] = None) -> List[IngestionBlock]:
        """
        Process a book file directly (bypassing queue system for testing).
        
        Args:
            file_path: Path to book file
            book_id: Book identifier
            metadata: Optional metadata dictionary
            
        Returns:
            List of ingestion blocks
        """
        message = {
            "book_id": book_id,
            "file_path": str(file_path),
            "original_filename": Path(file_path).name,
            "title": metadata.get("title", "Unknown") if metadata else "Unknown",
            "author": metadata.get("author", "Unknown") if metadata else "Unknown"
        }
        
        result = self.process(message)
        
        if result and result["status"] == "success":
            # Return the blocks from temporary storage
            blocks = self._temp_blocks
            self.logger.info(f"Direct processing completed for {book_id}")
            return blocks
        else:
            self.logger.error(f"Direct processing failed for {book_id}")
            return []


def main():
    """
    Main function for running TextIngester as standalone service.
    """
    # Setup logging
    logger = setup_logging("TextIngester")
    logger.info("🚀 Starting DIAS Stage A - Text Ingester")

    config = get_config()
    redis_client = get_redis_client(logger=logger)
    
    ingester = TextIngester(redis_client, config)
    
    print("TextIngester service starting...")
    print(f"Listening on queue: {ingester.input_queue}")
    print(f"Output to queue: {ingester.output_queue}")
    
    try:
        ingester.run()
    except KeyboardInterrupt:
        print("\nShutting down TextIngester...")
        ingester.stop()
    except Exception as e:
        logger.error(f"Fatal error: {e}")
        ingester.stop()


if __name__ == "__main__":
    main()