"""
Test fixtures for TextIngester testing.
Creates sample PDF, EPUB, and DOCX files for integration testing.
"""

import os
import tempfile
from pathlib import Path
from typing import List

# PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

# DOCX generation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# EPUB generation (simplified)
import zipfile
import xml.etree.ElementTree as ET


def create_sample_pdf(output_path: str, num_chapters: int = 3, paragraphs_per_chapter: int = 10) -> str:
    """
    Create a sample PDF file with structured content for testing.
    
    Args:
        output_path: Path where to save the PDF
        num_chapters: Number of chapters to create
        paragraphs_per_chapter: Paragraphs per chapter
        
    Returns:
        Path to created PDF file
    """
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=TA_CENTER
    )
    
    chapter_style = ParagraphStyle(
        'CustomChapter',
        parent=styles['Heading2'],
        fontSize=18,
        spaceAfter=20,
        spaceBefore=20
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=12,
        alignment=TA_JUSTIFY
    )
    
    story = []
    
    # Title page
    story.append(Paragraph("Sample Book for DIAS Testing", title_style))
    story.append(Spacer(1, 20))
    story.append(Paragraph("Generated for TextIngester Integration Tests", styles['Normal']))
    story.append(PageBreak())
    
    # Generate sample content
    sample_sentences = [
        "The art of digital storytelling has evolved dramatically in recent years.",
        "Modern technology enables us to create immersive narrative experiences.",
        "This transformation affects how we consume and interact with literature.",
        "Artificial intelligence plays a crucial role in content generation.",
        "The intersection of technology and creativity opens new possibilities.",
        "Digital formats allow for dynamic and interactive storytelling.",
        "Machine learning algorithms can analyze narrative structures.",
        "Content adaptation becomes more sophisticated with AI assistance.",
        "The future of publishing lies in intelligent automation.",
        "Human creativity remains essential in the digital age."
    ]
    
    for chapter_num in range(1, num_chapters + 1):
        story.append(Paragraph(f"Chapter {chapter_num}", chapter_style))
        
        for para_num in range(paragraphs_per_chapter):
            # Create paragraph with multiple sentences
            paragraph_text = " ".join([
                f"{sentence} This is additional content for paragraph {para_num + 1} in chapter {chapter_num}."
                for sentence in sample_sentences[:5]  # Use first 5 sentences
            ])
            
            story.append(Paragraph(paragraph_text, body_style))
            story.append(Spacer(1, 6))
        
        if chapter_num < num_chapters:
            story.append(PageBreak())
    
    # Build the PDF
    doc.build(story)
    return output_path


def create_sample_docx(output_path: str, num_chapters: int = 3, paragraphs_per_chapter: int = 10) -> str:
    """
    Create a sample DOCX file with structured content for testing.
    
    Args:
        output_path: Path where to save the DOCX
        num_chapters: Number of chapters to create
        paragraphs_per_chapter: Paragraphs per chapter
        
    Returns:
        Path to created DOCX file
    """
    doc = Document()
    
    # Title page
    title = doc.add_heading('Sample Book for DIAS Testing', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Generated for TextIngester Integration Tests')
    doc.add_page_break()
    
    # Sample content
    sample_sentences = [
        "The evolution of digital content creation has been remarkable.",
        "Technology enables new forms of narrative expression.",
        "Interactive storytelling creates engaging experiences.",
        "AI assists in content analysis and generation.",
        "Digital publishing transforms traditional workflows.",
        "Machine learning enhances creative processes.",
        "Automated systems improve content quality.",
        "Intelligent algorithms optimize text processing.",
        "Modern tools enable sophisticated content creation.",
        "The future combines human creativity with AI assistance."
    ]
    
    for chapter_num in range(1, num_chapters + 1):
        # Chapter heading
        chapter_heading = doc.add_heading(f'Chapter {chapter_num}', 1)
        
        for para_num in range(paragraphs_per_chapter):
            # Create paragraph
            paragraph_text = " ".join([
                f"{sentence} Additional content for paragraph {para_num + 1} in chapter {chapter_num}."
                for sentence in sample_sentences[:5]
            ])
            
            paragraph = doc.add_paragraph(paragraph_text)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Add some spacing
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(6)
    
    # Save the document
    doc.save(output_path)
    return output_path


def create_sample_epub(output_path: str, num_chapters: int = 3, paragraphs_per_chapter: int = 10) -> str:
    """
    Create a simplified sample EPUB file for testing.
    Note: This creates a basic EPUB structure, not a full-featured one.
    
    Args:
        output_path: Path where to save the EPUB
        num_chapters: Number of chapters to create
        paragraphs_per_chapter: Paragraphs per chapter
        
    Returns:
        Path to created EPUB file
    """
    
    # Create temporary directory for EPUB contents
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        
        # Create mimetype file (required for EPUB)
        (temp_path / "mimetype").write_text("application/epub+zip")
        
        # Create META-INF directory
        meta_inf = temp_path / "META-INF"
        meta_inf.mkdir()
        
        # Create container.xml
        container_xml = """<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
    <rootfiles>
        <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
    </rootfiles>
</container>"""
        (meta_inf / "container.xml").write_text(container_xml)
        
        # Create OEBPS directory
        oebps = temp_path / "OEBPS"
        oebps.mkdir()
        
        # Create content.opf (package document)
        content_opf = f"""<?xml version="1.0" encoding="UTF-8"?>
<package version="2.0" xmlns="http://www.idpf.org/2007/opf">
    <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
        <dc:title>Sample Book for DIAS Testing</dc:title>
        <dc:creator>Test Author</dc:creator>
        <dc:language>en</dc:language>
        <dc:identifier id="bookid">sample-book-123</dc:identifier>
    </metadata>
    <manifest>
        <item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>
        <item id="content" href="content.html" media-type="application/xhtml+xml"/>
        <item id="css" href="styles.css" media-type="text/css"/>
    </manifest>
    <spine toc="ncx">
        <itemref idref="content"/>
    </spine>
</package>"""
        (oebps / "content.opf").write_text(content_opf)
        
        # Create toc.ncx (navigation document)
        toc_ncx = """<?xml version="1.0" encoding="UTF-8"?>
<ncx version="2005-1" xmlns="http://www.daisy.org/z3986/2005/ncx/">
    <head>
        <meta name="dtb:uid" content="sample-book-123"/>
        <meta name="dtb:depth" content="1"/>
        <meta name="dtb:totalPageCount" content="0"/>
        <meta name="dtb:maxPageNumber" content="0"/>
    </head>
    <docTitle>
        <text>Sample Book for DIAS Testing</text>
    </docTitle>
    <navMap>
        <navPoint id="content" playOrder="1">
            <navLabel>
                <text>Content</text>
            </navLabel>
            <content src="content.html"/>
        </navPoint>
    </navMap>
</ncx>"""
        (oebps / "toc.ncx").write_text(toc_ncx)
        
        # Create styles.css
        styles_css = """body { font-family: serif; margin: 2em; line-height: 1.6; }
h1 { color: #333; border-bottom: 2px solid #333; }
h2 { color: #666; margin-top: 2em; }
p { text-align: justify; margin-bottom: 1em; }"""
        (oebps / "styles.css").write_text(styles_css)
        
        # Create content.html
        sample_sentences = [
            "Digital storytelling has transformed modern literature.",
            "Technology enables immersive narrative experiences.",
            "AI assists in content creation and analysis.",
            "Interactive formats engage readers in new ways.",
            "Machine learning enhances creative workflows.",
            "Digital publishing reaches global audiences.",
            "Automated systems improve content quality.",
            "Intelligent algorithms optimize processing.",
            "Modern tools enable sophisticated creation.",
            "Innovation drives the future of publishing."
        ]
        
        html_content = """<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>Sample Book for DIAS Testing</title>
    <link rel="stylesheet" type="text/css" href="styles.css"/>
</head>
<body>
    <h1>Sample Book for DIAS Testing</h1>
    <p>Generated for TextIngester Integration Tests</p>
"""
        
        for chapter_num in range(1, num_chapters + 1):
            html_content += f"    <h2>Chapter {chapter_num}</h2>\n"
            
            for para_num in range(paragraphs_per_chapter):
                paragraph_text = " ".join([
                    f"{sentence} Additional content for paragraph {para_num + 1} in chapter {chapter_num}."
                    for sentence in sample_sentences[:5]
                ])
                html_content += f"    <p>{paragraph_text}</p>\n"
        
        html_content += """</body>
</html>"""
        (oebps / "content.html").write_text(html_content)
        
        # Create the EPUB ZIP file
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as epub_zip:
            # Add mimetype first (must be uncompressed)
            epub_zip.write(temp_path / "mimetype", "mimetype", compress_type=zipfile.ZIP_STORED)
            
            # Add other files
            for root, dirs, files in os.walk(temp_path):
                for file in files:
                    if file == "mimetype":
                        continue
                    
                    file_path = Path(root) / file
                    arc_path = file_path.relative_to(temp_path)
                    epub_zip.write(file_path, arc_path)
    
    return output_path


def create_test_files(output_dir: str) -> dict:
    """
    Create all test file types in the specified directory.
    
    Args:
        output_dir: Directory to create test files
        
    Returns:
        Dictionary with paths to created files
    """
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    files_created = {}
    
    # Create PDF
    pdf_path = output_path / "sample_book.pdf"
    create_sample_pdf(str(pdf_path))
    files_created['pdf'] = str(pdf_path)
    
    # Create DOCX
    docx_path = output_path / "sample_book.docx"
    create_sample_docx(str(docx_path))
    files_created['docx'] = str(docx_path)
    
    # Create EPUB
    epub_path = output_path / "sample_book.epub"
    create_sample_epub(str(epub_path))
    files_created['epub'] = str(epub_path)
    
    return files_created


if __name__ == "__main__":
    # Create test files in current directory
    test_files = create_test_files("tests/fixtures")
    print("Created test files:")
    for format_type, path in test_files.items():
        print(f"  {format_type.upper()}: {path}")
        print(f"    Size: {Path(path).stat().st_size} bytes")